"""Convert uploaded Word documents to PDF for report storage."""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

from django.core.files.base import ContentFile


def docx_upload_to_pdf_file(uploaded_file) -> ContentFile:
    pdf_bytes = convert_docx_to_pdf_bytes(uploaded_file)
    base = Path(uploaded_file.name or "report").stem
    return ContentFile(pdf_bytes, name=f"{base}.pdf")


def convert_docx_to_pdf_bytes(uploaded_file) -> bytes:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            return _convert_with_libreoffice(soffice, uploaded_file)
        except (FileNotFoundError, subprocess.CalledProcessError, OSError, TimeoutError):
            uploaded_file.seek(0)

    return _convert_docx_with_reportlab(uploaded_file)


def _convert_with_libreoffice(soffice: str, uploaded_file) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "document.docx"
        docx_path.write_bytes(uploaded_file.read())
        uploaded_file.seek(0)

        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_path),
                str(docx_path),
            ],
            check=True,
            timeout=120,
            capture_output=True,
        )

        pdf_path = tmp_path / "document.pdf"
        if not pdf_path.is_file():
            raise FileNotFoundError("LibreOffice did not produce a PDF file.")
        return pdf_path.read_bytes()


def _convert_docx_with_reportlab(uploaded_file) -> bytes:
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    uploaded_file.seek(0)
    doc = Document(uploaded_file)
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4, title="Report")
    styles = getSampleStyleSheet()
    story = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        safe = (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        story.append(Paragraph(safe, styles["Normal"]))
        story.append(Spacer(1, 6))

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                line = " · ".join(cells)
                safe = (
                    line.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                story.append(Paragraph(safe, styles["Normal"]))
                story.append(Spacer(1, 4))

    if not story:
        raise ValueError("The Word document has no readable text.")

    pdf.build(story)
    return buffer.getvalue()
